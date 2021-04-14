import docx
import random
import os
import re

dictionary = {}
test_file_on = ""

def give_absolute_text(text):
    text = re.sub("\d+", "", text)
    text = text.strip()
    pattern = '[:,-,)]'
    init_val = 1
    text = re.sub("-+", "-", re.sub(pattern, "-", text))
    if text[0] != '-':
        init_val = 0
    output = list(map(lambda x: x.strip(), text[init_val:].split('-')))
    return output


def getDictionary(filename):
    doc = docx.Document(filename)
    fullTextDict = {}
    for para in doc.paragraphs:
        if para.text != "":
            absolute_line_text = give_absolute_text(para.text)
            fullTextDict[absolute_line_text[1]] = absolute_line_text[0]
    return fullTextDict

def getDictionaryInv(filename):
    doc = docx.Document(filename)
    fullTextDict = {}
    absolute_line_text = None
    for para in doc.paragraphs:
        if para.text != "":
            absolute_line_text = give_absolute_text(para.text)
    for i in range(0, len(absolute_line_text), 2):
        fullTextDict[absolute_line_text[i]] = absolute_line_text[i+1]
    return fullTextDict
    
def create_test_from_dictionary(lim=None):
    key_vals_of_dictionary = list(dictionary.keys())
    random.shuffle(key_vals_of_dictionary)
    length_of_dictionary = len(dictionary)
    if lim == None or lim > length_of_dictionary:
        lim = length_of_dictionary
    return "\n".join([str(word_cnt + 1) + ') ' + key_vals_of_dictionary[word_cnt] + '-'for word_cnt in range(lim)])



def create_test(sourcefile, testfile=None, lim=None):
    if testfile == None:
        testfile = "test.docx"
    global dictionary, test_file_on
    dictionary = getDictionary(sourcefile)
    test_data = create_test_from_dictionary(lim)
    doc = docx.Document()
    doc.add_paragraph(test_data)
    doc.save(testfile)
    os.startfile(testfile)
    test_file_on = testfile

def check_test():
    test_dict = getDictionaryInv(test_file_on)
    mistakes = []
    score = 0
    tot = len(test_dict)
    for test_data in test_dict:
        if test_dict[test_data] == dictionary[test_data]:
            score += 1
        elif "/" in dictionary[test_data]:
            words = list(map(lambda x: x.strip(), dictionary[test_data].split("/")))
            for i in words:
                if test_dict[test_data] == i:
                    score += 1
                    break
            else:
                mistakes.append([test_data, dictionary[test_data], test_dict[test_data]])
            
        else:
            mistakes.append([test_data, dictionary[test_data], test_dict[test_data]])
    return (score, tot, (score / tot) * 100, mistakes)

    
